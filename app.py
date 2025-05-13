import streamlit as st
import os
from google import genai
from google.genai import types
from io import BytesIO
import mimetypes
from docx import Document

# Lấy API Key từ Streamlit Secrets
API_KEY = st.secrets["general"]["GENAI_API_KEY"]

# Kiểm tra API Key
if not API_KEY:
    raise ValueError("API Key is missing. Please set the GENAI_API_KEY in Streamlit secrets.")

# Hàm trích xuất nội dung từ file âm thanh
def generate_transcription(uploaded_file):
    # Khởi tạo client Google GenAI
    client = genai.Client(api_key=API_KEY)
    
    # Xác định MIME type từ tên tệp (hoặc bạn có thể xác định nó thủ công)
    mime_type, _ = mimetypes.guess_type(uploaded_file.name)
    
    if not mime_type:
        raise ValueError(f"Unable to guess MIME type for the file: {uploaded_file.name}")
    
    # Tải tệp lên Google GenAI bằng BytesIO
    audio_data = BytesIO(uploaded_file.getvalue())  # Sử dụng BytesIO để tạo đối tượng file-like
    
    # Tải tệp lên API GenAI
    files = [
        client.files.upload(file=audio_data, mime_type=mime_type),  # Cung cấp MIME type chính xác
    ]
    
    model = "gemini-2.5-flash-preview-04-17"
    
    contents = [
        types.Content(
            role="user",
            parts=[types.Part.from_uri(file_uri=files[0].uri, mime_type=files[0].mime_type)],
        ),
        types.Content(
            role="model",
            parts=[types.Part.from_text(text="lấy nội dung đàm thoại của file âm thanh này")],
        ),
    ]
    
    response = client.models.generate_content(model=model, contents=contents)
    return response.text

# Hàm tạo và tải xuống file Word (.docx)
def create_word_document(transcription):
    doc = Document()
    doc.add_heading('Transcription from Audio File', 0)
    doc.add_paragraph(transcription)
    
    # Lưu vào bộ nhớ tạm
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Giao diện Streamlit
def main():
    st.title("Audio Transcription to Word")
    st.write("Tải lên file âm thanh để trích xuất nội dung và lưu dưới dạng file Word.")
    
    # Upload file âm thanh
    uploaded_file = st.file_uploader("Chọn file âm thanh", type=["mp3", "m4a", "wav"])

    if uploaded_file is not None:
        # Gọi hàm để trích xuất nội dung
        transcription = generate_transcription(uploaded_file)
        st.write("Nội dung trích xuất từ âm thanh:")
        st.text_area("Transcript", transcription, height=200)

        # Tạo và cho phép tải xuống file Word
        word_file = create_word_document(transcription)
        st.download_button(
            label="Tải về file Word",
            data=word_file,
            file_name="transcription.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
